import streamlit as st
from pdf2image import convert_from_bytes
from PIL import Image
import easyocr
from docx import Document
from docx.shared import Inches
import numpy as np
import io

reader = easyocr.Reader(['ko'])

st.title("PDF → DOCX 이미지 변환기 (Streamlit)")

uploaded_pdf = st.file_uploader("PDF 파일을 업로드하세요", type=["pdf"])

def split_vertical(img):
    w, h = img.size
    left = img.crop((0, 0, w // 2, h))
    right = img.crop((w // 2, 0, w, h))
    return [left, right]

def detect_bracket_regions(image, keyword_prefix='[', keyword_suffix=']'):
    results = reader.readtext(np.array(image), detail=1)
    regions = []
    for res in results:
        text = res[1]
        if text.strip().startswith(keyword_prefix) and text.strip().endswith(keyword_suffix):
            regions.append(res[0])  # bounding box
    return regions

def crop_regions_by_brackets(img, regions):
    cropped_imgs = []
    h = img.height
    regions_sorted = sorted(regions, key=lambda x: x[0][1])  # sort by y
    y_coords = [0] + [int(box[0][1]) for box in regions_sorted] + [h]

    for i in range(len(y_coords)-1):
        cropped_imgs.append(img.crop((0, y_coords[i], img.width, y_coords[i+1])))
    return cropped_imgs

if uploaded_pdf:
    pages = convert_from_bytes(uploaded_pdf.read())
    doc = Document()

    for page_num, page in enumerate(pages):
        st.write(f"처리 중: 페이지 {page_num+1}")
        halves = split_vertical(page)

        for idx, half in enumerate(halves):
            bracket_regions = detect_bracket_regions(half)

            if bracket_regions:
                split_images = crop_regions_by_brackets(half, bracket_regions)
            else:
                split_images = [half]

            for img in split_images:
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)

                doc.add_picture(img_bytes, width=Inches(6))
                doc.add_paragraph("")  # 여백 추가

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    st.download_button(
        label="📥 변환된 문서 다운로드 (DOCX)",
        data=doc_bytes,
        file_name="converted.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
